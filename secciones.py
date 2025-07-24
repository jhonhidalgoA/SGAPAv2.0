# ─────────────────────────────────────────────────────
# ✅ ORDEN DE LAS ROUTES

# 1.  MATRÍCULA ESTUDIANTES
# 2.  REGISTRO DOCENTES
# 3.  CIRCULARES
# 4.  REGISTRAR CALIFICACIONES
# 5.  MENÚ ESCOLAR
# 6.  CALENDARIO
# 7.  PLANEACIÓN
# 8.  REPORTES
# 9.  ASISTENCIA
# ─────────────────────────────────────────────────────



# ─────────────────────────────────────────────────────
# IMPORTACIONES ESTÁNDAR
# ─────────────────────────────────────────────────────
import os
import re
import json
from datetime import datetime
from io import BytesIO
from functools import reduce


# ─────────────────────────────────────────────────────
# FLASK Y EXTENSIONES
# ─────────────────────────────────────────────────────
from flask import (
    Blueprint, render_template, request, redirect, url_for,
    flash, jsonify, current_app, make_response, send_file
)
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash

# ─────────────────────────────────────────────────────
# BASE DE DATOS (MySQL / Conexiones)
# ─────────────────────────────────────────────────────
import pymysql
from db import get_connection as get_db

# ─────────────────────────────────────────────────────
# GENERACIÓN DE DOCUMENTOS (PDF / Word)
# ─────────────────────────────────────────────────────
# ReportLab (PDF)
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle,
    Paragraph, Spacer, Image
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.pdfgen import canvas

# python-docx (Word)
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# ─────────────────────────────────────────────────────
# UTILIDADES PERSONALIZADAS
# ─────────────────────────────────────────────────────
from utils.consecutivo import obtener_siguiente_consecutivo



secciones = Blueprint('secciones', __name__, url_prefix='/secciones')

ORDEN_DIAS = ["Lunes", "Martes", "Miércoles", "Jueves", "Viernes"]
ORDEN_CATEGORIAS = [
    "Menú del día",
    "Menú vegetariano",
    "Menú alternativo",    
    "Complementos"
]




# ─────────────────────────────────────────────────────
# MATRÍCULA ESTUDIANTES  
# ─────────────────────────────────────────────────────
@secciones.route('/matricula', methods=['GET', 'POST'])
def matricula():
    return render_template('secciones/matricula.html')

@secciones.route('/guardar_matricula', methods=['POST'])
def guardar_matricula():
    db = get_db()
    cursor = db.cursor()

    # Campos obligatorios del estudiante
    nombres = request.form.get('student_name')
    apellidos = request.form.get('student_lastname')
    tipo_documento = request.form.get('student_document_type')
    numero_documento = request.form.get('student_document_number')
    email = request.form.get('student_email')

    if not nombres or not apellidos or not tipo_documento or not numero_documento or not email:
        flash("❌ Todos los campos obligatorios deben estar llenos", "danger")
        
        # Enviar campos vacíos para resaltarlos en HTML
        campos_vacios = []
        if not nombres: campos_vacios.append('student_name')
        if not apellidos: campos_vacios.append('student_lastname')
        if not tipo_documento: campos_vacios.append('student_document_type')
        if not numero_documento: campos_vacios.append('student_document_number')
        if not email: campos_vacios.append('student_email')

        return redirect(url_for('secciones.matricula', campos_vacios=','.join(campos_vacios)))
    
    if email and not re.match(r'^[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+$', email):
        flash("❌ El correo electrónico no es válido", "danger")
        return redirect(url_for('secciones.matricula'))


    # Verificar si ya existe el documento
    cursor.execute("SELECT student_id FROM estudiantes WHERE document_number = %s", (numero_documento,))
    estudiante_existente = cursor.fetchone()

    if estudiante_existente:
        flash("⚠️ Este estudiante ya está matriculado", "warning")
        return redirect(url_for('secciones.matricula'))

    # Más datos personales del estudiante
    full_name = f"{nombres} {apellidos}"
    fecha_nacimiento = request.form.get('student_birth_date')
    genero = request.form.get('student_gender')
    lugar_nacimiento = request.form.get('student_birth_place')
    telefono = request.form.get('student_phone')
    email = request.form.get('student_email')
    direccion = request.form.get('student_address')
    barrio = request.form.get('student_neighborhood')
    localidad = request.form.get('student_locality')
    estrato = request.form.get('student_socioeconomic_status')
    zona = request.form.get('student_zone')
    grupo_sanguineo = request.form.get('student_blood_group')
    eps = request.form.get('student_eps')
    discapacidad = 1 if request.form.get('student_disability') == 'Sí' else 0
    tipo_discapacidad = request.form.get('student_disability_type') or None
    toma_medicamentos = 1 if request.form.get('student_takes_medicine') == 'Sí' else 0
    tipo_medicamento = request.form.get('student_medicine_type') or None

    # Fecha de matrícula
    fecha_matricula = request.form.get('registration_date') or datetime.now().date()

    # Subir foto del estudiante
    photo = request.files.get('student_photo')
    photo_filename = None

    if photo and photo.filename != '':
        try:
            photo_filename = secure_filename(photo.filename)
            upload_path = os.path.join(current_app.config['UPLOAD_FOLDER'], photo_filename)
            photo.save(upload_path)
        except Exception as e:
            flash(f"❌ Error al guardar la imagen: {str(e)}", "danger")
            return redirect(url_for('secciones.matricula'))
    else:
        flash("⚠️ Debes subir una foto del estudiante", "warning")
        return redirect(url_for('secciones.matricula', campos_vacios='student_photo'))

    # Buscar grade_id desde grados
    grado_ingreso = request.form.get('student_grade')
    cursor.execute("SELECT grade_id FROM grados WHERE name = %s", (grado_ingreso,))
    grado_result = cursor.fetchone()

    if not grado_result:
        flash("⚠️ Grado no encontrado", "warning")
        return redirect(url_for('secciones.matricula'))

    grade_id = grado_result['grade_id']
    grupo = request.form.get('student_group')
    jornada = request.form.get('student_shift')

    # Insertar en tabla estudiantes
    try:
        cursor.execute("""
            INSERT INTO estudiantes (
                full_name, birth_date, gender, birth_place,
                document_type, document_number, phone, email,
                address, neighborhood, locality, socioeconomic_status,
                zone, blood_group, eps, disability,
                disability_type, takes_medicine, medicine_type, photo_path, enrollment_date
            ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (
            full_name, fecha_nacimiento, genero, lugar_nacimiento,
            tipo_documento, numero_documento, telefono, email,
            direccion, barrio, localidad, estrato,
            zona, grupo_sanguineo, eps, discapacidad,
            tipo_discapacidad, toma_medicamentos, tipo_medicamento,
            photo_filename, fecha_matricula
        ))
        db.commit()
        student_id = cursor.lastrowid

        # Crear usuario en users
        password_hash = generate_password_hash(numero_documento)
        cursor.execute("""
            INSERT INTO users (
                full_name, role_id, username, password_hash, document_number, email
            ) VALUES (%s, %s, %s, %s, %s, %s)
        """, (full_name, 4, numero_documento, password_hash, numero_documento, email))
        db.commit()
        user_id = cursor.lastrowid

        # Actualizar estudiante con user_id
        cursor.execute("UPDATE estudiantes SET user_id = %s WHERE student_id = %s", (user_id, student_id))
        db.commit()

        # Datos madre
        nombre_madre = request.form.get('mother_name')
        tipo_documento_madre = request.form.get('mother_document_type')
        documento_madre = request.form.get('mother_document_number')
        telefono_madre = request.form.get('mother_phone')
        email_madre = request.form.get('mother_email')
        profesion_madre = request.form.get('mother_profession')
        ocupacion_madre = request.form.get('mother_occupation')
        is_guardian_madre = 1 if request.form.get('mother_is_guardian') else 0

        if nombre_madre and documento_madre:
            cursor.execute("""
                INSERT INTO contactos_familiares (
                    student_id, full_name, document_type, document_number,
                    phone, email, profession, occupation, relationship, is_guardian
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                student_id, nombre_madre, tipo_documento_madre, documento_madre,
                telefono_madre, email_madre, profesion_madre, ocupacion_madre,
                'madre', is_guardian_madre
            ))
            db.commit()
            mother_id = cursor.lastrowid

            # Actualizar guardian_id en estudiantes
            cursor.execute("UPDATE estudiantes SET guardian_id = %s WHERE student_id = %s", (mother_id, student_id))
            db.commit()

        # Datos padre
        nombre_padre = request.form.get('father_name')
        tipo_documento_padre = request.form.get('father_document_type')
        documento_padre = request.form.get('father_document_number')
        telefono_padre = request.form.get('father_phone')
        email_padre = request.form.get('father_email')
        profesion_padre = request.form.get('father_profession')
        ocupacion_padre = request.form.get('father_occupation')
        is_guardian_padre = 1 if request.form.get('father_is_guardian') else 0

        if nombre_padre and documento_padre:
            cursor.execute("""
                INSERT INTO contactos_familiares (
                    student_id, full_name, document_type, document_number,
                    phone, email, profession, occupation, relationship, is_guardian
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                student_id, nombre_padre, tipo_documento_padre, documento_padre,
                telefono_padre, email_padre, profesion_padre, ocupacion_padre,
                'padre', is_guardian_padre
            ))
            db.commit()

        # Insertar en matricula_estudiantes
        cursor.execute("""
            INSERT INTO matricula_estudiantes (
                student_id, grade_id, group_name, shift, academic_year
            ) VALUES (%s, %s, %s, %s, %s)
        """, (student_id, grade_id, grupo, jornada, datetime.now().year))
        db.commit()

        # Mensaje final de éxito
        flash("✅ Estudiante matriculado correctamente", "success")
        return redirect(url_for('secciones.matricula'))

    except Exception as e:
        db.rollback()
        flash(f"❌ Error al guardar: {str(e)}", "danger")
        return redirect(url_for('secciones.matricula'))


@secciones.route('/editar_estudiante/<int:student_id>', methods=['GET', 'POST'])
def editar_estudiante(student_id):
    """
    Ruta para editar un estudiante específico
    """
    db = get_db()
    cursor = db.cursor(pymysql.cursors.DictCursor)
    
    if request.method == 'GET':
        try:
            # Obtener datos del estudiante
            cursor.execute("""
                SELECT 
                    e.*,
                    g.name as grado_nombre,
                    me.group_name,
                    me.shift,
                    me.academic_year
                FROM estudiantes e
                LEFT JOIN matricula_estudiantes me ON e.student_id = me.student_id
                LEFT JOIN grados g ON me.grade_id = g.grade_id
                WHERE e.student_id = %s
            """, (student_id,))
            
            estudiante = cursor.fetchone()
            
            if not estudiante:
                flash("❌ Estudiante no encontrado", "danger")
                return redirect(url_for('secciones.matricula'))
            
            # Obtener contactos familiares
            cursor.execute("""
                SELECT * FROM contactos_familiares 
                WHERE student_id = %s
            """, (student_id,))
            
            contactos = cursor.fetchall()
            
            # Obtener lista de grados para el select
            cursor.execute("SELECT * FROM grados ORDER BY name")
            grados = cursor.fetchall()
            
            return render_template('secciones/editar_estudiante.html', 
                                 estudiante=estudiante, 
                                 contactos=contactos, 
                                 grados=grados)
                                 
        except Exception as e:
            flash(f"❌ Error al cargar estudiante: {str(e)}", "danger")
            return redirect(url_for('secciones.matricula'))
    
    elif request.method == 'POST':
        try:
            # Actualizar datos del estudiante
            nombres = request.form.get('student_name')
            apellidos = request.form.get('student_lastname')
            full_name = f"{nombres} {apellidos}"
            
            # Validaciones básicas
            if not nombres or not apellidos:
                flash("❌ Nombres y apellidos son obligatorios", "danger")
                return redirect(url_for('secciones.editar_estudiante', student_id=student_id))
            
            # Actualizar estudiante
            cursor.execute("""
                UPDATE estudiantes SET
                    full_name = %s,
                    birth_date = %s,
                    gender = %s,
                    birth_place = %s,
                    document_type = %s,
                    document_number = %s,
                    phone = %s,
                    email = %s,
                    address = %s,
                    neighborhood = %s,
                    locality = %s,
                    socioeconomic_status = %s,
                    zone = %s,
                    blood_group = %s,
                    eps = %s
                WHERE student_id = %s
            """, (
                full_name,
                request.form.get('student_birth_date'),
                request.form.get('student_gender'),
                request.form.get('student_birth_place'),
                request.form.get('student_document_type'),
                request.form.get('student_document_number'),
                request.form.get('student_phone'),
                request.form.get('student_email'),
                request.form.get('student_address'),
                request.form.get('student_neighborhood'),
                request.form.get('student_locality'),
                request.form.get('student_socioeconomic_status'),
                request.form.get('student_zone'),
                request.form.get('student_blood_group'),
                request.form.get('student_eps'),
                student_id
            ))
            
            # Actualizar matrícula si hay cambios de grado
            grado_nuevo = request.form.get('student_grade')
            if grado_nuevo:
                cursor.execute("SELECT grade_id FROM grados WHERE name = %s", (grado_nuevo,))
                grado_result = cursor.fetchone()
                
                if grado_result:
                    cursor.execute("""
                        UPDATE matricula_estudiantes SET
                            grade_id = %s,
                            group_name = %s,
                            shift = %s
                        WHERE student_id = %s
                    """, (
                        grado_result['grade_id'],
                        request.form.get('student_group'),
                        request.form.get('student_shift'),
                        student_id
                    ))
            
            db.commit()
            flash("✅ Estudiante actualizado correctamente", "success")
            return redirect(url_for('secciones.matricula'))
            
        except Exception as e:
            db.rollback()
            flash(f"❌ Error al actualizar: {str(e)}", "danger")
            return redirect(url_for('secciones.editar_estudiante', student_id=student_id))
    
        finally:
            cursor.close()
            db.close()



@secciones.route('/limpiar_datos')
def limpiar_datos():
    db = get_db()
    cursor = db.cursor()

    try:
        # Desactivar restricciones de llaves foráneas
        cursor.execute("SET FOREIGN_KEY_CHECKS = 0;")        
        cursor.execute("DELETE FROM matricula_estudiantes;")
        cursor.execute("DELETE FROM contactos_familiares;")
        cursor.execute("DELETE FROM users;")
        cursor.execute("DELETE FROM estudiantes;")
        cursor.execute("DELETE FROM asistencias;")

        # Reiniciar contadores AUTO_INCREMENT
        cursor.execute("ALTER TABLE estudiantes AUTO_INCREMENT = 1;")
        cursor.execute("ALTER TABLE contactos_familiares AUTO_INCREMENT = 1;")
        cursor.execute("ALTER TABLE users AUTO_INCREMENT = 1;")
        cursor.execute("ALTER TABLE matricula_estudiantes AUTO_INCREMENT = 1;")
        cursor.execute("ALTER TABLE asistencias AUTO_INCREMENT = 1;")
        # Reactivar restricciones
        cursor.execute("SET FOREIGN_KEY_CHECKS = 1;")
        db.commit()

        flash("✅ Datos borrados correctamente", "success")
    except Exception as e:
        db.rollback()
        flash(f"❌ Error al limpiar datos: {str(e)}", "danger")
    
    return redirect(url_for('secciones.matricula'))


# Registrar Docente

@secciones.route('/registro_docente', methods=['GET', 'POST'])
def registro_docente():
    return render_template('secciones/docente_registrar.html')

@secciones.route('/guardar_docente', methods=['POST'])
def guardar_docente():
    db = get_db()
    cursor = db.cursor()

    try:
        # Obtener datos del formulario
        nombres = request.form.get('teacher_name')
        apellidos = request.form.get('teacher_lastname')
        fecha_registro = request.form.get('registration_date_teacher')
        codigo = request.form.get('codigo_teacher')
        fecha_nacimiento = request.form.get('teacher_birth_date')
        edad = request.form.get('teacher_age')
        genero = request.form.get('teacher_gender')
        lugar_nacimiento = request.form.get('teacher_birth_place')
        tipo_documento = request.form.get('teacher_document_type')
        numero_documento = request.form.get('teacher_document_number')
        telefono = request.form.get('teacher_phone')
        correo = request.form.get('teacher_email')
        profesion = request.form.get('profession')
        area = request.form.get('area')       
        resolucion = request.form.get('resolucion')
        escalafon = request.form.get('scale')  
        foto = request.files.get('student_photo')

        # Validación básica de campos obligatorios
        if not all([nombres, apellidos, tipo_documento, numero_documento, area, escalafon]):
            flash("❌ Todos los campos obligatorios deben estar llenos", "danger")
            return redirect(url_for('secciones.registro_docente'))

        # Validar número de documento (solo números)
        if not numero_documento.isdigit():
            flash("❌ El número de documento debe contener solo números", "danger")
            return redirect(url_for('secciones.registro_docente'))

        # Validar correo si se proporciona
        if correo and '@' not in correo:
            flash("❌ El correo electrónico no es válido", "danger")
            return redirect(url_for('secciones.registro_docente'))

        # Validar teléfono si se llena
        if telefono and (not telefono.isdigit() or len(telefono) < 7):
            flash("❌ El teléfono debe tener al menos 7 dígitos", "danger")
            return redirect(url_for('secciones.registro_docente'))

        # Verificar si ya existe el docente
        cursor.execute("SELECT teacher_id FROM docentes_datos WHERE document_number = %s", (numero_documento,))
        docente_existe = cursor.fetchone()
        if docente_existe:
            flash("⚠️ Este docente ya está registrado", "warning")
            return redirect(url_for('secciones.registro_docente'))

        # Subir foto del docente
        photo_path = None
        if foto and foto.filename != '':
            filename = secure_filename(foto.filename)
            upload_path = os.path.join('.', 'static', 'uploads', filename)
            foto.save(upload_path)
            photo_path = filename
        else:
            flash("⚠️ Debes subir una foto del docente", "warning")
            return redirect(url_for('secciones.registro_docente'))

        # Generar contraseña hash
        password_hash = generate_password_hash(numero_documento)

        # Insertar en tabla users
        cursor.execute("""
            INSERT INTO users (
                full_name, role_id, username, password_hash, document_number, email
            ) VALUES (%s, %s, %s, %s, %s, %s)
        """, (f"{nombres} {apellidos}", 2, numero_documento, password_hash, numero_documento, correo))
        db.commit()
        user_id = cursor.lastrowid

        # Insertar en tabla docentes_datos
        cursor.execute("""
            INSERT INTO docentes_datos (
                first_name, last_name, registration_date, code,
                birth_date, age, gender, birth_place,
                document_type, document_number, phone, email,
                profession, subject_area, resolution_number, scale, photo_path, user_id
            ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (
            nombres, apellidos, fecha_registro, codigo,
            fecha_nacimiento, edad, genero, lugar_nacimiento,
            tipo_documento, numero_documento, telefono, correo,
            profesion, area, resolucion, escalafon, photo_path, user_id
        ))
        db.commit()

        # Mensaje de éxito
        flash("✅ Docente matriculado correctamente", "success")

    except Exception as e:
        db.rollback()
        print(f"Error al registrar docente: {e}")
        flash(f"❌ Error al registrar docente: {str(e)}", "danger")
    finally:
        cursor.close()
        db.close()

    return redirect(url_for('secciones.registro_docente'))


@secciones.route('/api/docentes')
def get_docentes():
    try:
        conn = get_db()
        cursor = conn.cursor(pymysql.cursors.DictCursor)
        cursor.execute("""
            SELECT 
                teacher_id AS id,
                first_name AS teacher_name,
                last_name AS teacher_lastname,
                document_type AS teacher_document_type,
                document_number AS teacher_document_number,
                phone AS teacher_phone,
                email AS teacher_email,
                profession,
                subject_area AS area,
                resolution_number AS resolucion,
                scale,
                birth_date AS teacher_birth_date,
                age AS teacher_age,
                gender AS teacher_gender,
                birth_place AS teacher_birth_place,
                code AS codigo_teacher,
                photo_path
            FROM docentes_datos
        """)
        docentes = cursor.fetchall()
        return jsonify(docentes)
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()

@secciones.route('/actualizar_docente/<int:id>', methods=['POST'])
def actualizar_docente(id):
    try:
        # Obtener datos del formulario
        first_name = request.form['teacher_name']
        last_name = request.form['teacher_lastname']
        document_type = request.form['teacher_document_type']
        document_number = request.form['teacher_document_number']
        phone = request.form['teacher_phone']
        email = request.form['teacher_email']
        profession = request.form['profession']
        subject_area = request.form['area']
        resolution_number = request.form.get('resolucion')
        scale = request.form['scale']

        registration_date = request.form.get('registration_date_teacher')
        code = request.form.get('codigo_teacher')
        birth_date = request.form.get('teacher_birth_date')
        age = request.form.get('teacher_age')
        gender = request.form.get('teacher_gender')
        birth_place = request.form.get('teacher_birth_place')

        # Manejo opcional de foto
        photo = request.files.get('student_photo')
        upload_folder = current_app.config['UPLOAD_FOLDER']

        # Obtener foto actual desde DB
        conn = get_db()
        cursor = conn.cursor(pymysql.cursors.DictCursor)
        cursor.execute("SELECT photo_path FROM docentes_datos WHERE teacher_id = %s", (id,))
        existing = cursor.fetchone()
        photo_path = existing['photo_path'] if existing else None

        # Si hay nueva foto, guardarla y actualizar ruta
        if photo and photo.filename.strip() != '':
            filename = secure_filename(photo.filename)
            photo.save(os.path.join(upload_folder, filename))
            photo_path = f"uploads/{filename}"

        # Actualizar datos del docente
        cursor.execute("""
            UPDATE docentes_datos SET 
                first_name = %s, last_name = %s, document_type = %s,
                document_number = %s, phone = %s, email = %s,
                profession = %s, subject_area = %s, resolution_number = %s,
                scale = %s, code = %s, birth_date = %s, age = %s,
                gender = %s, birth_place = %s
            WHERE teacher_id = %s
        """, (
            first_name, last_name, document_type, document_number,
            phone, email, profession, subject_area, resolution_number,
            scale, code, birth_date, age, gender, birth_place, id
        ))

        # Solo actualizar foto si se carga una nueva
        if photo_path:
            cursor.execute("UPDATE docentes_datos SET photo_path = %s WHERE teacher_id = %s", (photo_path, id))

        conn.commit()

        flash("Docente actualizado correctamente", "success")
        return redirect(url_for('secciones.registro_docente'))

    except Exception as e:
        conn.rollback()
        flash(f"Error al actualizar el docente: {str(e)}", "danger")
        return redirect(url_for('secciones.registro_docente'))
    finally:
        cursor.close()
        conn.close()

# ─────────────────────────────────────────────────────
# CIRCULARES
# ─────────────────────────────────────────────────────
@secciones.route('/circulares', methods=['GET', 'POST'])
def circulares():    
    base_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(base_dir, 'static', 'data', 'circulares.json')

    if not os.path.exists(file_path):
        return "Archivo no encontrado", 500

    with open(file_path, encoding='utf-8') as f:
        lista_circulares = json.load(f)

    return render_template('secciones/circulares.html', circulares=lista_circulares)  


# ─────────────────────────────────────────────────────
# REGISTRAR CALIFICACIONES
# ─────────────────────────────────────────────────────
@secciones.route('/calificaciones', methods=['GET', 'POST'])
def calificaciones():
    return render_template('secciones/calificaciones.html')

# REEMPLAZAR la función api_estudiantes con esta versión corregida

@secciones.route('/api/estudiantes', methods=['GET'])
@secciones.route('/api/estudiantes/<string:grupo>', methods=['GET'])
def api_estudiantes(grupo=None):
    """
    API endpoint con formato de nombres y grados corregido
    """
    db = get_db()
    cursor = db.cursor(pymysql.cursors.DictCursor)
    
    try:
        # Obtener parámetro de búsqueda si existe
        search = request.args.get('search', '').strip()
        print(f"🔍 Búsqueda: '{search}', Grupo: '{grupo}'")
        
        # Query según los parámetros (mismo código anterior)
        if grupo:
            if search:
                query = """
                    SELECT 
                        e.student_id,
                        e.full_name,
                        e.document_number,
                        e.email,
                        e.phone,
                        g.name as grado_nombre,
                        m.group_name,
                        m.shift
                    FROM matricula_estudiantes m
                    JOIN estudiantes e ON m.student_id = e.student_id
                    LEFT JOIN grados g ON m.grade_id = g.grade_id
                    WHERE m.grade_id = %s
                      AND m.academic_year = YEAR(CURRENT_DATE)
                      AND (e.full_name LIKE %s OR e.document_number LIKE %s)
                    ORDER BY e.full_name ASC
                """
                search_param = f"%{search}%"
                cursor.execute(query, (grupo, search_param, search_param))
            else:
                query = """
                    SELECT 
                        e.student_id,
                        e.full_name,
                        e.document_number,
                        e.email,
                        e.phone,
                        g.name as grado_nombre,
                        m.group_name,
                        m.shift
                    FROM matricula_estudiantes m
                    JOIN estudiantes e ON m.student_id = e.student_id
                    LEFT JOIN grados g ON m.grade_id = g.grade_id
                    WHERE m.grade_id = %s
                      AND m.academic_year = YEAR(CURRENT_DATE)
                    ORDER BY e.full_name ASC
                """
                cursor.execute(query, (grupo,))
        else:
            if search:
                query = """
                    SELECT 
                        e.student_id,
                        e.full_name,
                        e.document_number,
                        e.email,
                        e.phone,
                        g.name as grado_nombre,
                        m.group_name,
                        m.shift
                    FROM estudiantes e
                    LEFT JOIN matricula_estudiantes m ON e.student_id = m.student_id 
                        AND m.academic_year = YEAR(CURRENT_DATE)
                    LEFT JOIN grados g ON m.grade_id = g.grade_id
                    WHERE (e.full_name LIKE %s 
                           OR e.document_number LIKE %s 
                           OR g.name LIKE %s
                           OR m.group_name LIKE %s)
                    ORDER BY e.full_name ASC
                    LIMIT 100
                """
                search_param = f"%{search}%"
                cursor.execute(query, (search_param, search_param, search_param, search_param))
            else:
                query = """
                    SELECT 
                        e.student_id,
                        e.full_name,
                        e.document_number,
                        e.email,
                        e.phone,
                        g.name as grado_nombre,
                        m.group_name,
                        m.shift
                    FROM estudiantes e
                    LEFT JOIN matricula_estudiantes m ON e.student_id = m.student_id 
                        AND m.academic_year = YEAR(CURRENT_DATE)
                    LEFT JOIN grados g ON m.grade_id = g.grade_id
                    ORDER BY e.full_name ASC
                    LIMIT 100
                """
                cursor.execute(query)

        estudiantes = cursor.fetchall()
        print(f"📊 Estudiantes encontrados: {len(estudiantes)}")

        if not estudiantes:
            message = f"No se encontraron estudiantes para el grupo {grupo}" if grupo else "No se encontraron estudiantes"
            return jsonify({
                "success": True,
                "data": [],
                "total": 0,
                "message": message
            })

        # FUNCIÓN CORREGIDA para separar nombres y apellidos
        def separar_nombres_apellidos_correcto(nombre_completo):
            """
            Función que separa correctamente nombres de apellidos
            Ejemplos:
            - "Andrea del Pilar Cardona Morales" → nombres: "Andrea del Pilar", apellidos: "Cardona Morales"
            - "Daniela Cárdenas López" → nombres: "Daniela", apellidos: "Cárdenas López"
            - "Juan David Torres Mejía" → nombres: "Juan David", apellidos: "Torres Mejía"
            """
            if not nombre_completo:
                return "", ""
            
            # Limpiar espacios extra
            partes = [parte.strip() for parte in nombre_completo.strip().split() if parte.strip()]
            total_partes = len(partes)
            
            if total_partes == 0:
                return "", ""
            elif total_partes == 1:
                return partes[0], ""
            elif total_partes == 2:
                # Solo 2 palabras: asumimos nombre + apellido
                return partes[0], partes[1]
            elif total_partes == 3:
                # 3 palabras: primer nombre + dos apellidos
                return partes[0], f"{partes[1]} {partes[2]}"
            elif total_partes == 4:
                # 4 palabras: dos nombres + dos apellidos
                return f"{partes[0]} {partes[1]}", f"{partes[2]} {partes[3]}"
            elif total_partes == 5:
                # 5 palabras: tres nombres + dos apellidos (ej: Andrea del Pilar Cardona Morales)
                return f"{partes[0]} {partes[1]} {partes[2]}", f"{partes[3]} {partes[4]}"
            else:
                # 6+ palabras: mitad nombres, mitad apellidos
                mitad = total_partes // 2
                nombres = " ".join(partes[:mitad])
                apellidos = " ".join(partes[mitad:])
                return nombres, apellidos

        # FUNCIÓN para formatear grados con guión
        def formatear_grado_con_guion(grado_nombre, group_name):
            """
            Formatea grados como "septimo - 701" o "sexto - 601"
            """
            if not grado_nombre:
                return "Sin asignar"
            
            # Mantener el nombre del grado original (septimo, sexto, etc.)
            grado_formateado = grado_nombre.lower()
            
            # Si tiene group_name, agregarlo con guión y espacios
            if group_name:
                return f"{grado_formateado} - {group_name}"
            else:
                return grado_formateado

        # Procesar datos para el frontend
        data = []
        for est in estudiantes:
            # USAR LA FUNCIÓN CORREGIDA para nombres y apellidos
            nombres, apellidos = separar_nombres_apellidos_correcto(est["full_name"])
            
            # USAR LA FUNCIÓN para formatear grados con guión
            grado_completo = formatear_grado_con_guion(est.get('grado_nombre'), est.get('group_name'))

            data.append({
                "id": est["student_id"],
                "student_id": str(est["student_id"]),
                "full_name": est["full_name"],
                "nombres": nombres,
                "apellidos": apellidos,
                "documento": est.get("document_number"),
                "email": est.get("email"),
                "telefono": est.get("phone"),
                "grado": grado_completo,
                "jornada": est.get("shift")
            })

        # Log para debugging
        print("📋 Ejemplos de formato:")
        for i, item in enumerate(data[:3]):  # Mostrar solo los primeros 3
            print(f"  {i+1}. '{item['full_name']}' → nombres: '{item['nombres']}', apellidos: '{item['apellidos']}', grado: '{item['grado']}'")

        return jsonify({
            "success": True,
            "data": data,
            "total": len(data),
            "message": f"Se encontraron {len(data)} estudiantes"
        })

    except Exception as e:
        print(f"❌ Error al cargar estudiantes: {str(e)}")
        import traceback
        traceback.print_exc()
        
        return jsonify({
            "success": False,
            "data": [],
            "total": 0,
            "error": "No se pudieron cargar los estudiantes",
            "message": f"Error interno: {str(e)}"
        }), 500
    
    finally:
        if cursor:
            cursor.close()
        if db:
            db.close()



@secciones.route('/api/estudiante/<int:student_id>', methods=['GET'])
def obtener_estudiante_detalle(student_id):
    """
    Obtener detalles completos de un estudiante específico
    """
    db = get_db()
    cursor = db.cursor(pymysql.cursors.DictCursor)
    
    try:
        # Obtener estudiante con todos sus datos
        cursor.execute("""
            SELECT 
                e.*,
                g.name as grado_nombre,
                m.group_name,
                m.shift,
                m.academic_year
            FROM estudiantes e
            LEFT JOIN matricula_estudiantes m ON e.student_id = m.student_id
                AND m.academic_year = YEAR(CURRENT_DATE)
            LEFT JOIN grados g ON m.grade_id = g.grade_id
            WHERE e.student_id = %s
        """, (student_id,))
        
        estudiante = cursor.fetchone()
        
        if not estudiante:
            return jsonify({
                'success': False,
                'message': 'Estudiante no encontrado'
            }), 404
        
        # Obtener contactos familiares si existen
        cursor.execute("""
            SELECT * FROM contactos_familiares 
            WHERE student_id = %s
            ORDER BY relationship, full_name
        """, (student_id,))
        
        contactos = cursor.fetchall()
        
        return jsonify({
            'success': True,
            'data': {
                'estudiante': estudiante,
                'contactos_familiares': contactos
            },
            'message': 'Estudiante encontrado correctamente'
        })
        
    except Exception as e:
        print(f"❌ Error al obtener estudiante {student_id}: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'Error al obtener estudiante: {str(e)}'
        }), 500
    
    finally:
        if cursor:
            cursor.close()
        if db:
            db.close()

# TAMBIÉN AGREGAR RUTA PARA EDITAR (si quieres esa funcionalidad)
@secciones.route('/editar_estudiante/<int:student_id>', methods=['GET'])
def editar_estudiante_form(student_id):
    """
    Mostrar formulario de edición de estudiante
    """
    # Aquí renderizarías un template para editar el estudiante
    # Por ahora solo un mensaje de confirmación
    return f"<h1>Editar Estudiante ID: {student_id}</h1><p>Aquí iría el formulario de edición</p>"

# ENDPOINT DE TEST PARA VERIFICAR QUE TODO FUNCIONA
@secciones.route('/api/test', methods=['GET'])
def test_endpoint():
    """
    Endpoint de prueba
    """
    return jsonify({
        'success': True,
        'message': 'API funcionando correctamente',
        'timestamp': datetime.now().isoformat(),
        'endpoints_disponibles': [
            '/api/estudiantes - Obtener todos los estudiantes',
            '/api/estudiantes/<grupo> - Obtener estudiantes de un grupo',
            '/api/estudiante/<id> - Obtener detalles de un estudiante',
            '/api/test - Este endpoint de prueba'
        ]
    })

# Guardar nuevas notas
@secciones.route('/api/guardar-notas', methods=['POST'])
def guardar_notas():
    data = request.get_json()

    if not data or not isinstance(data, dict):
        return jsonify({
            "error": "Datos inválidos",
            "details": "No se recibió un JSON válido"
        }), 400

    grupo = data.get('grupo')
    asignatura = data.get('asignatura')
    periodo = data.get('periodo')
    notas_data = data.get('notas', [])

    if not all([grupo, asignatura, periodo]):
        return jsonify({
            "error": "Faltan parámetros obligatorios",
            "missing": {
                "grupo": grupo is None,
                "asignatura": asignatura is None,
                "periodo": periodo is None
            }
        }), 400

    if not isinstance(notas_data, list) or len(notas_data) == 0:
        return jsonify({
            "error": "No hay estudiantes con notas para guardar",
            "details": "El campo 'notas' está vacío o no es una lista"
        }), 400

    db = None
    cursor = None

    try:
        db = get_db()
        cursor = db.cursor()

        # Eliminar notas previas del grupo, asignatura y período
        cursor.execute("""
            DELETE FROM valor_notas 
            WHERE grade_id = %s AND subject_id = %s AND period_id = %s
        """, (grupo, asignatura, periodo))

        # Insertar nuevas notas
        for nota in notas_data:
            student_id = nota.get("student_id")
            nota1 = float(nota.get("nota1")) if nota.get("nota1") else None
            nota2 = float(nota.get("nota2")) if nota.get("nota2") else None
            nota3 = float(nota.get("nota3")) if nota.get("nota3") else None
            nota4 = float(nota.get("nota4")) if nota.get("nota4") else None
            nota5 = float(nota.get("nota5")) if nota.get("nota5") else None

            cursor.execute("""
                INSERT INTO valor_notas (
                    grade_id, student_id, subject_id, period_id, 
                    nota1, nota2, nota3, nota4, nota5
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                grupo, student_id, asignatura, periodo,
                nota1, nota2, nota3, nota4, nota5
            ))

        db.commit()

        return jsonify({
            "message": "✅ Notas guardadas correctamente",
            "total_registros": len(notas_data),
            "grupo": grupo,
            "asignatura": asignatura,
            "periodo": periodo
        })

    except Exception as e:
        if db:
            db.rollback()
        print(f"❌ Error al guardar notas: {str(e)}")
        return jsonify({
            "error": "No se pudieron guardar las notas",
            "details": str(e)
        }), 500

    finally:
        if cursor:
            cursor.close()
        if db and hasattr(db, 'close'):
            db.close()

# Cargar notas existentes
@secciones.route('/api/cargar-notas', methods=['GET'])
def cargar_notas():
    grupo = request.args.get('grupo')
    asignatura = request.args.get('asignatura')
    periodo = request.args.get('periodo')

    if not all([grupo, asignatura, periodo]):
        return jsonify({
            "error": "Faltan parámetros obligatorios",
            "missing": {
                "grupo": grupo is None,
                "asignatura": asignatura is None,
                "periodo": periodo is None
            }
        }), 400

    try:
        db = get_db()
        cursor = db.cursor(pymysql.cursors.DictCursor)

        cursor.execute("""
            SELECT student_id, nota1, nota2, nota3, nota4, nota5
            FROM valor_notas
            WHERE grade_id = %s AND subject_id = %s AND period_id = %s
        """, (grupo, asignatura, periodo))

        notas = cursor.fetchall()
        return jsonify({"data": notas})

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    
@secciones.route('/api/notas-estudiante', methods=['GET'])
def get_notas_estudiante():
    grupo = request.args.get('grupo')
    estudiante_id = request.args.get('estudiante')
    periodo = request.args.get('periodo') or '1'  

    if not all([grupo, estudiante_id]):
        return jsonify({"error": "Faltan parámetros"}), 400

    try:
        db = get_db()
        cursor = db.cursor(pymysql.cursors.DictCursor)

        # Obtener datos del estudiante
        cursor.execute("""
            SELECT e.full_name, e.document_number, m.grade_id 
            FROM estudiantes e
            JOIN matricula_estudiantes m ON e.student_id = m.student_id
            WHERE e.student_id = %s AND m.grade_id = %s
        """, (estudiante_id, grupo))
        estudiante = cursor.fetchone()
        
        # Obtener notas del estudiante
        cursor.execute("""
            SELECT v.nota1, v.nota2, v.nota3, v.nota4, v.nota5, s.name AS asignatura
            FROM valor_notas v
            JOIN asignaturas s ON v.subject_id = s.subject_id
            WHERE v.grade_id = %s AND v.student_id = %s AND v.period_id = %s
        """, (grupo, estudiante_id, periodo))
        notas = cursor.fetchall()

        return jsonify({
            "estudiante": estudiante,
            "notas": notas
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500    
    
@secciones.route('/api/asignaturas', methods=['GET'])
def get_asignaturas():
    try:
        db = get_db()
        cursor = db.cursor(pymysql.cursors.DictCursor)
        cursor.execute("SELECT subject_id AS id, name AS nombre FROM asignaturas ORDER BY name ASC")
        asignaturas = cursor.fetchall()
        return jsonify({"data": asignaturas})
    except Exception as e:
        return jsonify({"error": str(e)}), 500    
    

# ─────────────────────────────────────────────────────
# MENU ESCOLAR
# ─────────────────────────────────────────────────────
def reordenar_menu(menu):
    ordered = {}
    for dia in ORDEN_DIAS:
        if dia in menu:
            ordered[dia] = {}

            # Ordenar categorías dentro del día
            for categoria in ORDEN_CATEGORIAS:
                if categoria in menu[dia]:
                    ordered[dia][categoria] = menu[dia][categoria]
    return ordered

@secciones.route('/guardar-cambios', methods=['POST'])
def guardar_cambios():
    data = request.get_json()

    dia = data['dia']
    categoria = data['categoria']
    original = data['platoOriginal']
    nuevo = data['platoNuevo']

    ruta_json = os.path.join(current_app.static_folder, "data", "menu.json")

    # Leer menú actual
    with open(ruta_json, "r", encoding="utf-8") as f:
        menu = json.load(f)

    encontrado = False
    for i, plato in enumerate(menu[dia][categoria]):
        if plato['nombre'] == original['nombre'] and plato['img'] == original['img']:
            menu[dia][categoria][i] = { "nombre": nuevo['nombre'], "img": nuevo['img'] }
            encontrado = True
            break

    if not encontrado:
        return jsonify({ "success": False, "error": "Plato no encontrado" })

    # Reordenar todo antes de guardar
    menu_ordenado = reordenar_menu(menu)

    # Guardar menú actualizado
    with open(ruta_json, "w", encoding="utf-8") as f:
        json.dump(menu_ordenado, f, indent=4, ensure_ascii=False)

    return jsonify({ "success": True })

@secciones.route('/editar_menu', methods=['GET', 'POST'])
def editar_menu(): 
    ruta_json = os.path.join(current_app.static_folder, "data", "menu.json")
    with open(ruta_json, "r", encoding="utf-8") as archivo:
        menus = json.load(archivo) 
    menus_ordenados = reordenar_menu(menus)     
    return render_template("secciones/editar_menu.html", menus=menus_ordenados)

@secciones.route("/restaurante")
def restaurante():
    ruta_json = os.path.join(current_app.static_folder, "data", "menu.json")
    with open(ruta_json, "r", encoding="utf-8") as archivo:
        menus = json.load(archivo)
    menus_ordenados = reordenar_menu(menus)
    semana = "14 al 18 de julio"  
    return render_template("secciones/restaurante.html", menus=menus_ordenados, semana=semana)

def cargar_menu():
    ruta_json = os.path.join(current_app.static_folder, "data", "menu.json")
    with open(ruta_json, "r", encoding="utf-8") as f:
        menu = json.load(f)

    return reordenar_menu(menu)

# ─────────────────────────────────────────────────────
# CALENDARIO
# ─────────────────────────────────────────────────────
@secciones.route('/calendario')
def calendario():
    return render_template('secciones/calendario.html')

@secciones.route('/obtener-eventos', methods=['GET'])
def obtener_eventos():
    connection = None
    try:
        connection = get_db()  
        with connection.cursor() as cursor:
            sql = "SELECT id, titulo, fecha_inicio, fecha_fin, color FROM eventos_calendario"
            cursor.execute(sql)
            resultados = cursor.fetchall()

            eventos = []
            for row in resultados:
                evento = {
                    'id': row['id'],
                    'title': row['titulo'],
                    'startDate': row['fecha_inicio'].strftime('%Y-%m-%d'),
                    'endDate': row['fecha_fin'].strftime('%Y-%m-%d'),
                    'color': row['color']
                }
                eventos.append(evento)

            return jsonify({'success': True, 'events': eventos})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

    finally:
        if connection:
            connection.close()


@secciones.route('/guardar-evento', methods=['POST'])
def guardar_evento():
    data = request.get_json()
    evento_id = data.get('id')
    title = data.get('title')
    start_date = data.get('startDate')
    end_date = data.get('endDate')
    color = data.get('color')

    if not all([title, start_date, end_date, color]):
        return jsonify({
            'success': False,
            'error': 'Faltan campos obligatorios'
        }), 400

    connection = None
    try:
        connection = get_db()
        with connection.cursor() as cursor:
            sql = """
                INSERT INTO eventos_calendario (id, titulo, fecha_inicio, fecha_fin, color)
                VALUES (%s, %s, %s, %s, %s)
                ON DUPLICATE KEY UPDATE
                    titulo = VALUES(titulo),
                    fecha_inicio = VALUES(fecha_inicio),
                    fecha_fin = VALUES(fecha_fin),
                    color = VALUES(color)
            """
            cursor.execute(sql, (evento_id, title, start_date, end_date, color))
        connection.commit()
        return jsonify({'success': True})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

    finally:
        if connection:
            connection.close()


@secciones.route('/eliminar-evento', methods=['POST'])
def eliminar_evento():
    data = request.get_json()
    evento_id = data.get('id')

    if not evento_id:
        return jsonify({
            'success': False,
            'error': 'ID del evento no proporcionado'
        }), 400

    connection = None
    try:
        connection = get_db()
        with connection.cursor() as cursor:
            sql = "DELETE FROM eventos_calendario WHERE id = %s"
            cursor.execute(sql, (evento_id,))
        connection.commit()
        return jsonify({'success': True})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

    finally:
        if connection:
            connection.close()

@secciones.route('/ver-eventos', methods=['GET'])
def ver_eventos():
    connection = None
    try:
        connection = get_db()
        with connection.cursor() as cursor:
            sql = "SELECT id, titulo, fecha_inicio, fecha_fin, color FROM eventos_calendario"
            cursor.execute(sql)
            resultados = cursor.fetchall()

            eventos = []
            for row in resultados:
                evento = {
                    'id': row['id'],
                    'title': row['titulo'],
                    'start': row['fecha_inicio'].strftime('%Y-%m-%d'),
                    'end': row['fecha_fin'].strftime('%Y-%m-%d'),
                    'color': row['color']
                }
                eventos.append(evento)

            return jsonify({'success': True, 'events': eventos})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

    finally:
        if connection:
            connection.close()

@secciones.route('/ver-calendario')
def ver_calendario():
    return render_template('secciones/calendario_ver.html')

# ─────────────────────────────────────────────────────
# PLANEACIÓN
# ─────────────────────────────────────────────────────

@secciones.route('/generar_pdf', methods=['POST'])
def generar_pdf():
    datos = {
        'fecha_inicio': request.form.get('dateStar'),
        'fecha_fin': request.form.get('dateEnd'),
        'periodo': request.form.get('periodo'),
        'grado': request.form.get('grado'),
        'asignatura': request.form.get('asignatura'),
        'tipo': request.form.get('tipo'),
        'estandar': request.form.get('estandar'),
        'dba': request.form.get('dba'),
        'evidencia': request.form.get('evidencia'),
        'competencias': request.form.get('competencias'),
        'objetivos': request.form.get('objetivos'),
        'proyecto_transversal': request.form.get('proyecto_transversal')
    }

   
    if not all([datos['fecha_inicio'], datos['fecha_fin'], datos['periodo'], datos['grado'], datos['asignatura']]):
        flash("⚠️ Faltan campos obligatorios", "warning")
        return redirect(url_for('secciones.planeacion'))

    # Obtener siguiente consecutivo
    tipo_documento = 'plan-clase'    

    try:
        consecutivo = obtener_siguiente_consecutivo(tipo_documento)
    except Exception as e:
        flash(f"❌ Error al obtener el consecutivo: {str(e)}", "danger")
        return redirect(url_for('secciones.planeacion'))

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            rightMargin=40, leftMargin=40,
                            topMargin=40, bottomMargin=40)
    story = []

    # Estilos
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Encabezado', fontSize=12, alignment=1, spaceAfter=10))
    styles.add(ParagraphStyle(name='Titulo', fontSize=16, alignment=1, textColor=colors.darkblue, spaceAfter=20))
    styles.add(ParagraphStyle(name='Seccion', fontSize=12, textColor=colors.black, spaceBefore=14, spaceAfter=6))
    styles.add(ParagraphStyle(name='NormalEspaciado', fontSize=10, spaceAfter=10))

    # Logo (si luego tienes uno, reemplaza esta ruta)
    logo_path = 'static/img/logoColegio.png'
    if os.path.exists(logo_path):
       img = Image(logo_path, width=0.8*inch, height=0.8*inch)
       img.hAlign = 'LEFT'
       story.append(img)

    # Encabezado
    story.append(Paragraph("INSTITUCIÓN EDUCATIVA SGAPA", styles['Encabezado']))
    story.append(Paragraph("Docente: Jhon Fredy Hidalgo", styles['Encabezado']))
    story.append(Paragraph(f"PLAN DE CLASE {str(consecutivo).zfill(5)}", styles['Titulo']))
    story.append(Spacer(1, 12))

    # Información general en dos columnas
    info_data = [
        ['Fecha de Inicio:', datos['fecha_inicio'] or 'No especificada', 'Fecha de Fin:', datos['fecha_fin'] or 'No especificada'],
        ['Grado:', datos['grado'] or 'No especificado', 'Período:', datos['periodo'] or 'No especificado'],
        ['Asignatura:', datos['asignatura'] or 'No especificada', 'Tipo:', datos['tipo'] or 'No especificado']
    ]
    info_table = Table(info_data, colWidths=[1.5*inch, 2.5*inch, 1.5*inch, 2.5*inch])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('BACKGROUND', (2, 0), (2, -1), colors.lightgrey),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 10)
    ]))
    story.append(info_table)
    story.append(Spacer(1, 14))

    # Secciones de contenido
    def agregar_seccion(titulo, contenido):
        if contenido:
            story.append(Paragraph(titulo, styles['Seccion']))
            story.append(Paragraph(contenido, styles['NormalEspaciado']))

    agregar_seccion("ESTÁNDAR", datos['estandar'])
    agregar_seccion("DBA (Derechos Básicos de Aprendizaje)", datos['dba'])
    agregar_seccion("EVIDENCIAS DE APRENDIZAJE", datos['evidencia'])
    agregar_seccion("COMPETENCIAS", datos['competencias'])
    agregar_seccion("OBJETIVOS", datos['objetivos'])
    agregar_seccion("PROYECTO PEDAGÓGICO TRANSVERSAL", datos['proyecto_transversal'])

    # Fecha de generación
    story.append(Spacer(1, 24))
    fecha_generacion = datetime.now().strftime("%d/%m/%Y %H:%M")
    story.append(Paragraph(f"Documento generado el: {fecha_generacion}", styles['NormalEspaciado']))

    # Tabla de firmas
    story.append(Spacer(1, 40))
    firma_data = [
        ["____________________________", "____________________________"],
        ["Firma del Docente", "Firma del Coordinador o Rector"]
    ]
    firma_table = Table(firma_data, colWidths=[3*inch, 3*inch])
    firma_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 10)
    ]))
    story.append(firma_table)

    # Generar PDF
    doc.build(story)
    buffer.seek(0)
    response = make_response(buffer.getvalue())
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = f'attachment; filename=plan-clase-{str(consecutivo).zfill(5)}.pdf'
    return response

@secciones.route('/planeacion', methods=['GET', 'POST'])
def planeacion():
    db = get_db()
    cursor = db.cursor()
    tipo_documento = 'plan_clase'

    try:
        cursor.execute("""
            SELECT MAX(consecutivo) AS max_consecutivo 
            FROM pdf_consecutivos 
            WHERE tipo_documento = %s
        """, (tipo_documento,))
        
        resultado = cursor.fetchone()

        if resultado is None:
            ultimo_consecutivo = '00000'
        else:
            if isinstance(resultado, dict):
                valor = resultado.get('max_consecutivo')
            else:
                valor = resultado[0] if len(resultado) > 0 else None
            
            if valor is None:
                ultimo_consecutivo = '00000'
            else:
                ultimo_consecutivo = str(valor).zfill(5)

        return render_template('secciones/planeacion.html', ultimo_consecutivo=ultimo_consecutivo)

    except Exception as e:
        print(f"❌ Error al cargar consecutivo: {str(e)}")
        return render_template('secciones/planeacion.html', ultimo_consecutivo='00000')

    finally:
        cursor.close()
        db.close()

@secciones.route('/ver-consecutivos')
def ver_consecutivos():
    db = get_db()
    cursor = db.cursor()

    try:
        cursor.execute("SELECT * FROM pdf_consecutivos ORDER BY fecha_generacion DESC")
        registros = cursor.fetchall()
        
        return render_template('secciones/consecutivos.html', registros=registros)
    
    except Exception as e:
        flash(f"❌ Error al cargar consecutivos: {str(e)}", "danger")
        return redirect(url_for('secciones.index'))
    
    finally:
        cursor.close()
        db.close()

@secciones.route('/eliminar-consecutivo/<int:id>')
def eliminar_consecutivo(id):
    db = get_db()
    cursor = db.cursor()

    try:
        cursor.execute("DELETE FROM pdf_consecutivos WHERE id = %s", (id,))
        db.commit()
        flash("🗑️ Registro eliminado correctamente", "success")
    
    except Exception as e:
        db.rollback()
        flash(f"❌ No se pudo eliminar: {str(e)}", "danger")
    
    finally:
        cursor.close()
        db.close()
    
    return redirect(url_for('secciones.ver_consecutivos'))

@secciones.route('/editar-consecutivo/<int:id>', methods=['GET', 'POST'])
def editar_consecutivo(id):
    db = get_db()
    cursor = db.cursor()

    if request.method == 'POST':
        nuevo_consecutivo = request.form.get('consecutivo')
        tipo_documento = request.form.get('tipo_documento')

        if not nuevo_consecutivo.isdigit() or not tipo_documento:
            flash("⚠️ Datos inválidos", "warning")
            return redirect(url_for('secciones.editar_consecutivo', id=id))

        try:
            cursor.execute("""
                UPDATE pdf_consecutivos 
                SET consecutivo = %s, tipo_documento = %s 
                WHERE id = %s
            """, (nuevo_consecutivo, tipo_documento, id))
            db.commit()
            flash("✏️ Consecutivo actualizado", "success")
            return redirect(url_for('secciones.ver_consecutivos'))
        
        except Exception as e:
            db.rollback()
            flash(f"❌ Error al actualizar: {str(e)}", "danger")
            return redirect(url_for('secciones.ver_consecutivos'))

    try:
        cursor.execute("SELECT * FROM pdf_consecutivos WHERE id = %s", (id,))
        registro = cursor.fetchone()
        return render_template('secciones/editar_consecutivo.html', registro=registro)
    
    except Exception as e:
        flash(f"❌ Error al cargar datos: {str(e)}", "danger")
        return redirect(url_for('secciones.ver_consecutivos'))
    
    finally:
        cursor.close()
        db.close()

# ─────────────────────────────────────────────────────
# REPORTES
# ─────────────────────────────────────────────────────

@secciones.route('/reportes', methods=['GET', 'POST'])
def reportes():
    return render_template('secciones/reportes.html')

@secciones.route('/generar-certificado-word', methods=['POST'])
def generar_certificado_word():
    data = request.get_json()
    grupo = data.get('grupo')
    estudiante_id = data.get('estudiante_id')

    if not grupo or not estudiante_id:
        return jsonify({"error": "Faltan parámetros obligatorios"}), 400

    try:
        db = get_db()
        cursor = db.cursor(pymysql.cursors.DictCursor)

        # Obtener datos del estudiante
        cursor.execute("""
            SELECT 
                e.student_id,
                e.full_name,
                e.document_number,
                m.grade_id
            FROM estudiantes e
            JOIN matricula_estudiantes m ON e.student_id = m.student_id
            WHERE e.student_id = %s AND m.grade_id = %s
        """, (estudiante_id, grupo))
        estudiante = cursor.fetchone()

        if not estudiante:
            return jsonify({"error": "Estudiante no encontrado en este grado"}), 404

        # Cargar plantilla Word
        try:
            doc = Document("Colegio Educativo SGAPA.docx")
        except Exception as e:
            return jsonify({"error": "No se encontró la plantilla Word", "details": str(e)}), 500

        # Función para reemplazar texto con formato
        def replace_with_format(paragraph, placeholder, text, bold=False, uppercase=False):
            if placeholder in paragraph.text:
                before, after = paragraph.text.split(placeholder, 1)
                paragraph.clear()
                if before:
                    paragraph.add_run(before)
                    processed_text = text.upper() if uppercase else text
                run = paragraph.add_run(processed_text)
                run.bold = bold
                if after:
                    paragraph.add_run(after)

        # === Obtener fecha actual y formatearla ===
        fecha_actual = datetime.now()
        dia_actual = str(fecha_actual.day)

        MESES_ESPANOL = {
            "January": "Enero",
            "February": "Febrero",
            "March": "Marzo",
            "April": "Abril",
            "May": "Mayo",
            "June": "Junio",
            "July": "Julio",
            "August": "Agosto",
            "September": "Septiembre",
            "October": "Octubre",
            "November": "Noviembre",
            "December": "Diciembre"
        }

        mes_actual = fecha_actual.strftime("%B")  # Mes en inglés
        mes_actual_espanol = MESES_ESPANOL.get(mes_actual, mes_actual)  # Traducir a español
        año_actual = str(fecha_actual.year)

        # Reemplazar los placeholders
        for paragraph in doc.paragraphs:
            replace_with_format(paragraph, '{nombre}', estudiante['full_name'], bold=True, uppercase=True)
            replace_with_format(paragraph, '{documento}', str(estudiante['document_number']), bold=True, uppercase=True)
            replace_with_format(paragraph, '{grado}', str(grupo))
            replace_with_format(paragraph, '{jornada}', 'única')
            replace_with_format(paragraph, '{año}', año_actual)
            replace_with_format(paragraph, '{día}', dia_actual)
            replace_with_format(paragraph, '{mes}', mes_actual_espanol)
            replace_with_format(paragraph, '{año de fecha}', año_actual)

        # Guardar archivo temporal
        archivo_temporal = f"Certificado_{estudiante_id}.docx"
        doc.save(archivo_temporal)

        # Enviar archivo al cliente
        response = send_file(
            archivo_temporal,
            as_attachment=True,
            download_name=f"Certificado_{estudiante['full_name'].replace(' ', '_')}.docx"
        )

        # Eliminar el archivo temporal después de enviarlo
        try:
            os.remove(archivo_temporal)
        except Exception as e:
            print(f"[ERROR] No se pudo eliminar el archivo temporal: {e}")

        return response

    except Exception as e:
        print(f"[ERROR] {str(e)}")
        return jsonify({
            "error": "Error interno del servidor",
            "details": str(e)
        }), 500

@secciones.route('/generar-boletin-word', methods=['POST'])
def generar_boletin_word():
    data = request.get_json()
    grupo = data.get('grupo')
    estudiante_id = data.get('estudiante_id')
    periodo = data.get('periodo', '1')

    if not grupo or not estudiante_id:
        return jsonify({"error": "Faltan parámetros obligatorios"}), 400

    try:
        db = get_db()
        cursor = db.cursor(pymysql.cursors.DictCursor)

        # Obtener datos del estudiante
        cursor.execute("""
            SELECT 
                e.student_id,
                e.full_name,
                e.document_number,
                m.grade_id
            FROM estudiantes e
            JOIN matricula_estudiantes m ON e.student_id = m.student_id
            WHERE e.student_id = %s AND m.grade_id = %s
        """, (estudiante_id, grupo))
        estudiante = cursor.fetchone()

        if not estudiante:
            return jsonify({"error": "Estudiante no encontrado en este grado"}), 404

        # Cargar documento Word
        try:
            doc = Document("Boletin SGAPA.docx")
        except Exception as e:
            return jsonify({
                "error": "No se encontró la plantilla Word",
                "details": str(e)
            }), 500

        # Función para reemplazar texto con formato
        def replace_with_format(paragraph, placeholder, text, bold=False, uppercase=False):
            if placeholder in paragraph.text:
                before, after = paragraph.text.split(placeholder, 1)
                paragraph.clear()
                if before:
                    paragraph.add_run(before)
                processed_text = text.upper() if uppercase else text
                run = paragraph.add_run(processed_text)
                run.bold = bold
                if after:
                    paragraph.add_run(after)

        # === Obtener fecha actual ===
        fecha_actual = datetime.now()
        dia_actual = str(fecha_actual.day)

        MESES_ESPANOL = {
            "January": "Enero", "February": "Febrero", "March": "Marzo",
            "April": "Abril", "May": "Mayo", "June": "Junio",
            "July": "Julio", "August": "Agosto", "September": "Septiembre",
            "October": "Octubre", "November": "Noviembre", "December": "Diciembre"
        }

        mes_actual = fecha_actual.strftime("%B")
        mes_actual_espanol = MESES_ESPANOL.get(mes_actual, mes_actual)
        año_actual = str(fecha_actual.year)

        # Obtener notas finales
        cursor.execute("""
            SELECT v.nota_final, s.name AS asignatura
            FROM valor_notas v
            JOIN asignaturas s ON v.subject_id = s.subject_id
            WHERE v.grade_id = %s AND v.student_id = %s AND v.period_id = %s
        """, (grupo, estudiante_id, periodo))
        notas = cursor.fetchall()

        # Obtener intensidad horaria SOLO de las asignaturas con notas (CORREGIDO)
        cursor.execute("""
            SELECT DISTINCT s.name, s.hours_per_week 
            FROM asignaturas s
            JOIN valor_notas v ON s.subject_id = v.subject_id
            WHERE v.grade_id = %s AND v.student_id = %s AND v.period_id = %s
        """, (grupo, estudiante_id, periodo))
        intensidades = cursor.fetchall()

        # Función para limpiar nombres de asignaturas para placeholders
        def clean_name_for_placeholder(name):
            import re
            # Remover tildes y caracteres especiales, convertir a minúsculas
            name = name.lower()
            name = name.replace('á', 'a').replace('é', 'e').replace('í', 'i')
            name = name.replace('ó', 'o').replace('ú', 'u').replace('ñ', 'n')
            # Reemplazar espacios y caracteres especiales con guión bajo
            name = re.sub(r'[^a-z0-9]', '_', name)
            # Remover guiones bajos múltiples
            name = re.sub(r'_+', '_', name)
            # Remover guiones bajos al inicio y final
            name = name.strip('_')
            return name

        # Crear diccionarios de reemplazo
        nota_dict = {}
        intensidad_dict = {}

        # Mapear notas con nombres limpios
        for nota in notas:
            clean_name = clean_name_for_placeholder(nota['asignatura'])
            placeholder = f"{{nota_{clean_name}}}"
            nota_dict[placeholder] = str(nota['nota_final']) if nota['nota_final'] is not None else "0.0"

        # Mapear intensidades horarias con nombres limpios
        for intensidad in intensidades:
            clean_name = clean_name_for_placeholder(intensidad['name'])
            placeholder = f"{{intensidad_{clean_name}}}"
            intensidad_dict[placeholder] = str(intensidad['hours_per_week']) if intensidad['hours_per_week'] is not None else "0"

        # Crear diccionario completo de reemplazos
        all_replacements = {
            "{nombre}": estudiante['full_name'],
            "{documento}": str(estudiante['document_number']),
            "{grado}": str(grupo),
            "{periodo}": str(periodo),
            "{día}": dia_actual,
            "{mes}": mes_actual_espanol,
            "{año}": año_actual
        }

        # Agregar notas e intensidades
        all_replacements.update(nota_dict)
        all_replacements.update(intensidad_dict)

        # Reemplazar en párrafos (MÉTODO MEJORADO)
        for paragraph in doc.paragraphs:
            texto_actual = paragraph.text
            
            # Verificar si hay placeholders para reemplazar
            placeholders_encontrados = [p for p in all_replacements.keys() if p in texto_actual]
            
            if placeholders_encontrados:
                print(f"[DEBUG] Reemplazando en párrafo: '{texto_actual}'")
                
                # Reemplazar todos los placeholders en el texto
                for placeholder, value in all_replacements.items():
                    if placeholder in texto_actual:
                        texto_actual = texto_actual.replace(placeholder, str(value))
                
                # Limpiar el párrafo y agregar el texto modificado
                paragraph.clear()
                
                # Si contiene el nombre, aplicar formato especial
                if "{nombre}" in paragraph.text or "nombre" in texto_actual.lower():
                    # Buscar el nombre en el texto y aplicar formato
                    if estudiante['full_name'].upper() in texto_actual.upper():
                        partes = texto_actual.split(estudiante['full_name'])
                        if len(partes) > 1:
                            paragraph.add_run(partes[0])
                            run = paragraph.add_run(estudiante['full_name'].upper())
                            run.bold = True
                            paragraph.add_run(''.join(partes[1:]))
                        else:
                            paragraph.add_run(texto_actual)
                    else:
                        paragraph.add_run(texto_actual)
                else:
                    paragraph.add_run(texto_actual)
                
                print(f"[DEBUG] Resultado: '{paragraph.text}'")

        # Reemplazar en tablas (MÉTODO MEJORADO)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        texto_actual = paragraph.text
                        
                        # Verificar si hay placeholders para reemplazar
                        placeholders_encontrados = [p for p in all_replacements.keys() if p in texto_actual]
                        
                        if placeholders_encontrados:
                            # Reemplazar todos los placeholders en el texto
                            for placeholder, value in all_replacements.items():
                                if placeholder in texto_actual:
                                    texto_actual = texto_actual.replace(placeholder, str(value))
                            
                            # Limpiar el párrafo y agregar el texto modificado
                            paragraph.clear()
                            
                            # Si contiene el nombre, aplicar formato especial
                            if estudiante['full_name'].upper() in texto_actual.upper():
                                partes = texto_actual.split(estudiante['full_name'])
                                if len(partes) > 1:
                                    paragraph.add_run(partes[0])
                                    run = paragraph.add_run(estudiante['full_name'].upper())
                                    run.bold = True
                                    paragraph.add_run(''.join(partes[1:]))
                                else:
                                    paragraph.add_run(texto_actual)
                            else:
                                paragraph.add_run(texto_actual)

        # Guardar archivo temporal
        archivo_temporal = f"Boletin_{estudiante_id}_{periodo}.docx"
        doc.save(archivo_temporal)

        # Crear nombre de descarga limpio
        nombre_limpio = estudiante['full_name'].replace(' ', '_').replace('á', 'a').replace('é', 'e').replace('í', 'i').replace('ó', 'o').replace('ú', 'u')
        download_name = f"Boletin_{nombre_limpio}_Periodo_{periodo}.docx"

        # Enviar como descarga
        response = send_file(
            archivo_temporal,
            as_attachment=True,
            download_name=download_name
        )

        # Eliminar archivo temporal después de enviarlo
        try:
            os.remove(archivo_temporal)
        except Exception as e:
            print(f"[ERROR] No se pudo eliminar el archivo temporal: {e}")

        return response

    except Exception as e:
        print(f"[ERROR] Error en generar_boletin_word: {str(e)}")
        return jsonify({
            "error": "Error interno del servidor",
            "details": str(e)
        }), 500

# ─────────────────────────────────────────────────────
#  ASISTENCIA
# ─────────────────────────────────────────────────────

@secciones.route('/asistencia', methods=['GET', 'POST'])
def asistencia(): 
    return render_template('secciones/asistencia.html')

@secciones.route('/api/guardar-asistencia', methods=['POST'])
def guardar_asistencia():
    data = request.get_json()
    if not data:
        return jsonify({'success': False, 'message': 'No se recibieron datos'}), 400

    grado = data.get('grado')
    asignatura = data.get('asignatura')
    periodo = data.get('periodo')
    estudiantes = data.get('estudiantes')
    fecha = data.get('fecha')

    if not all([grado, asignatura, periodo, estudiantes, fecha]):
        return jsonify({'success': False, 'message': 'Datos incompletos'}), 400

    db = get_db()
    cursor = db.cursor()

    try:
        for est in estudiantes:
            student_id = est.get('student_id') or est.get('id')

            # 🔍 Validación: Comprobar si el estudiante existe
            cursor.execute("SELECT COUNT(*) AS count FROM estudiantes WHERE student_id = %s", (student_id,))
            result = cursor.fetchone()
            if not result or result['count'] == 0:
                return jsonify({'success': False, 'message': f'El estudiante {student_id} no existe'}), 400

            fallas = est.get('fallas', 0)
            observacion = est.get('observacion', '').strip()

            status = 'presente' if fallas == 0 else 'ausente'
            if observacion:
                status = 'justificado'

            query = """
                INSERT INTO asistencias (student_id, subject_id, date, status, notes)
                VALUES (%s, %s, %s, %s, %s)
                ON DUPLICATE KEY UPDATE
                    status = VALUES(status),
                    notes = VALUES(notes)
            """
            params = (student_id, asignatura, fecha, status, observacion)
            cursor.execute(query, params)

        db.commit()
        return jsonify({'success': True, 'message': 'Asistencia guardada correctamente'})

    except Exception as e:
        db.rollback()
        return jsonify({'success': False, 'message': f'Error al guardar: {str(e)}'}), 500


@secciones.route('/tareas', methods=['GET', 'POST'])
def tareas():
    return render_template('secciones/tareas.html')

# ─────────────────────────────────────────────────────
#  HORARIO
# ─────────────────────────────────────────────────────
@secciones.route('/horario', methods=['GET', 'POST'])
def horario():
    if request.method == 'POST':
        data = request.get_json()
        print("Datos recibidos:", data)  # Debug
        
        db = get_db()
        cursor = db.cursor()
        
        try:
            for entry in data:
                print(f"Procesando: {entry}")  # Debug
                
                # Verificar que la materia existe
                cursor.execute("SELECT subject_id FROM asignaturas WHERE name = %s", (entry['materia'],))
                subject_result = cursor.fetchone()
                
                if not subject_result:
                    print(f"Materia no encontrada: {entry['materia']}")
                    continue
                    
                cursor.execute("""
                    INSERT INTO horarios 
                    (schedule_day, start_time, end_time, subject_id, grade_id, period_id, teacher_id)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                """, (
                    entry['dia'],
                    entry['hora_inicio'],
                    entry['hora_fin'],
                    subject_result['subject_id'],
                    entry['grado_id'],
                    1,
                    entry['docente_id']
                ))
                
            db.commit()
            return jsonify({"status": "success", "message": "Horario guardado"}), 200
            
        except Exception as e:
            print(f"Error: {e}")  # Debug
            db.rollback()
            return jsonify({"status": "error", "message": str(e)}), 500    
    
    db = get_db()
    cursor = db.cursor()
    
    try:
        cursor.execute("SELECT user_id, full_name FROM users WHERE role_id = 2")
        docentes = cursor.fetchall()
    except Exception as e:
        print(f"Error al obtener docentes: {e}")
        docentes = []
    
    docentes_lista = [(d['user_id'], d['full_name']) for d in docentes] if docentes else []
    
    return render_template('secciones/horario.html', docentes=docentes_lista)

def format_timedelta(td):
    if isinstance(td, str):
        return td.split('.')[0]
    elif isinstance(td, timedelta):
        total_seconds = int(td.total_seconds())
        hours, remainder = divmod(total_seconds, 3600)
        minutes, _ = divmod(remainder, 60)
        return f"{hours:02d}:{minutes:02d}"
    else:
        return "00:00"


@secciones.route('/horario/ver/<int:grado_id>')
def ver_horario(grado_id):
    db = get_db()
    cursor = db.cursor()

    try:
        cursor.execute("""
            SELECT h.schedule_day AS dia, h.start_time AS hora_inicio, 
                   h.end_time AS hora_fin, a.name AS materia
            FROM horarios h
            JOIN asignaturas a ON h.subject_id = a.subject_id
            WHERE h.grade_id = %s
            ORDER BY FIELD(h.schedule_day, 'Lunes', 'Martes', 'Miércoles', 'Jueves', 'Viernes'),
                     h.start_time
        """, (grado_id,))
        
        horario = cursor.fetchall()

        # Convertir tiempos a formato HH:MM
        horario_formateado = []
        for row in horario:
            horario_formateado.append({
                "dia": row['dia'],
                "hora_inicio": format_timedelta(row['hora_inicio']),
                "hora_fin": format_timedelta(row['hora_fin']),
                "materia": row['materia']
            })

    except Exception as e:
        print("Error al cargar horario:", e)
        horario_formateado = []

    return render_template('secciones/ver_horario.html', grado_id=grado_id, horario=horario_formateado)


from datetime import timedelta
@secciones.route('/api/horario/<int:grade_id>', methods=['GET'])
def get_horario(grade_id):
    try:
        db = get_db()
        cursor = db.cursor(pymysql.cursors.DictCursor)

        # Alias claros para los campos de tiempo
        query = """
            SELECT 
                h.schedule_day AS dia,
                h.start_time AS hora_inicio,
                h.end_time AS hora_fin,
                a.name AS materia,
                h.teacher_id,
                h.grade_id
            FROM horarios h
            LEFT JOIN asignaturas a ON h.subject_id = a.subject_id
            WHERE h.grade_id = %s
            ORDER BY FIELD(h.schedule_day, 'Lunes', 'Martes', 'Miércoles', 'Jueves', 'Viernes'),
                     h.start_time
        """

        cursor.execute(query, (grade_id,))
        horario = cursor.fetchall()

        if not horario:
            return jsonify({"error": "No se encontró horario para este grado"}), 404

        # Formatear tiempos de start y end
        for fila in horario:
            for key in ['hora_inicio', 'hora_fin']:
                value = fila[key]
                if isinstance(value, timedelta):
                    total_seconds = int(value.total_seconds())
                    hours = total_seconds // 3600
                    minutes = (total_seconds % 3600) // 60
                    fila[key] = f"{hours:02d}:{minutes:02d}"
                elif isinstance(value, str):
                    fila[key] = value.split('.')[0]  # quitar microsegundos

        return jsonify(horario)

    except Exception as e:
        print(f"Error al obtener horario: {e}")
        return jsonify({
            "error": "Error interno del servidor",
            "detalle": str(e)
        }), 500


# Ruta para obtener materia por ID
@secciones.route('/secciones/api/materia/<int:subject_id>', methods=['GET'])
def get_materia(subject_id):
    db = get_db()
    cursor = db.cursor(pymysql.cursors.DictCursor)
    cursor.execute("SELECT * FROM asignaturas WHERE subject_id = %s", (subject_id,))
    materia = cursor.fetchone()
    return jsonify(materia)


@secciones.route('/borrar-horario/<int:grade_id>', methods=['POST'])
def borrar_horario(grade_id):
    db = get_db()
    cursor = db.cursor()

    try:
        cursor.execute("DELETE FROM horarios WHERE grade_id = %s", (grade_id,))
        db.commit()
        return jsonify({"success": True, "message": "Horario eliminado correctamente"}), 200
    except Exception as e:
        db.rollback()
        return jsonify({"success": False, "error": str(e)}), 500







@secciones.route('/usuarios', methods=['GET', 'POST'])
def usuarios():
    return render_template('secciones/usuarios.html')

@secciones.route('/eventos')
def eventos():
    return render_template('secciones/eventos.html')  

@secciones.route('/admisiones')
def admisiones():
    return render_template('secciones/admisiones.html')

@secciones.route('/noticias')
def noticias():
    return render_template('secciones/noticias.html')

@secciones.route('/ver-actividades')
def ver_actividades():
    return render_template('secciones/ver_actividades.html')


@secciones.route('/docente-generar-pdf-from-form', methods=['POST'])
def docente_generarPDF_from_form():
    # Obtener los datos enviados desde el formulario
    docente = {
        'registration_date_teacher': request.form.get('registration_date_teacher'),
        'first_name': request.form.get('first_name', ''),
        'last_name': request.form.get('last_name', ''),
        'document_number': request.form.get('document_number', ''),
        'email': request.form.get('email', ''),
        'phone': request.form.get('phone', ''),
        'profession': request.form.get('profession', ''),
        'subject_area': request.form.get('subject_area', ''),
        'age': request.form.get('age', ''),
        'gender': request.form.get('gender', ''),
        'document_type': request.form.get('document_type', ''),
        'birth_place': request.form.get('birth_place', ''),
        'birth_date': request.form.get('birth_date', ''),
        'code': request.form.get('code', ''),
        'resolution_number': request.form.get('resolution_number', ''),
        'scale': request.form.get('scale', ''),
        'photo_path': request.form.get('photo_path', '')
    }
    
    return render_template('secciones/docente_generarPDF.html', docente=docente)


@secciones.route('/docente-generar-pdf')
def docente_generarPDF():
    try:
        conn = get_db()
        cursor = conn.cursor(pymysql.cursors.DictCursor)        
        
        cursor.execute("""
            SELECT 
                first_name,
                last_name,
                document_number,
                email,
                profession,
                subject_area,
                phone,
                age,
                gender,
                document_type,
                birth_place,
                birth_date,
                code,
                resolution_number,
                scale,
                photo_path       
            FROM docentes_datos 
            ORDER BY teacher_id DESC 
            LIMIT 1
        """)
        
        docente = cursor.fetchone()        
        
        return render_template('secciones/docente_generarPDF.html', docente=docente)
        
    except Exception as e:
        print(f"Error al obtener datos del docente: {e}")
        return render_template('secciones/docente_generarPDF.html', docente=None)
    finally:
        cursor.close()
        conn.close()


@secciones.route('/ver_docente')
def ver_docente():
    return render_template('secciones/ver_docente.html')        